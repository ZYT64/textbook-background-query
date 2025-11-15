from flask import Flask, render_template, request, make_response, send_file
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn  # 关键导入：知乎文章核心函数
import io
from urllib.parse import quote
import time
import re

wait = []
app = Flask(__name__)

# app.py 开头新增：读取环境变量
import os
from dotenv import load_dotenv  # 新增：需要安装 python-dotenv
load_dotenv()  # 加载环境变量

# 修改AI调用函数中的api_key：从环境变量读取
def ask_AI(input_text, client_ip):
    api_key = os.getenv("AI_API_KEY")  # 不再硬编码！
    api_base = "https://open.bigmodel.cn/api/paas/v4/"
    client = OpenAI(api_key=api_key, base_url=api_base)
    try:
        response = client.chat.completions.create(
            model="glm-4-flash",
            messages=[{"role": "user", "content": input_text}],
            stream=False,
            temperature=0.7,
            max_tokens=4096,
            extra_headers={"lora_id": "0"},
            stream_options={"include_usage": True}
        )
        if client_ip in wait:
            wait.remove(client_ip)
        return response.choices[0].message.content
    except Exception as e:
        if client_ip in wait:
            wait.remove(client_ip)
        return f"AI调用失败：{str(e)}"

# 生成字号选项（不变）
def generate_font_sizes():
    font_sizes = []
    current = 3.0
    while current <= 24.0:
        font_sizes.append(round(current, 1))
        current += 0.5
    return font_sizes

# 核心修改：按知乎文章方法设置中文字体（兼容所有版本）
def generate_word_doc(textbook_title, selected_options, word_count, font_size, line_height, ai_response):
    doc = Document()
    target_font = "微软雅黑"  # 统一中文字体
    target_font_size = Pt(float(font_size))  # 统一字号
    target_line_spacing = float(line_height)  # 统一行间距

    # 1. 深度清理AI内容（不变）
    def clean_ai_content(content):
        content = re.sub(r"\*\*|\*|#|- |\+ |= |~", "", content)
        content = re.sub(r"\d+\.|\d+\)|①|②|③|④|⑤|⑴|⑵|⑶", "", content)
        content = re.sub(r"\[.*?\]|\(.*?\)|\{.*?\}|<.*?>|【.*?】|《.*?》", "", content)
        content = re.sub(r"[\x00-\x1F\x7F\x80-\x9F]", "", content)
        content = re.sub(r"\s+", " ", content).strip()
        content = re.sub(r"([。！？；：])", r"\1\n", content)
        return content

    cleaned_ai_response = clean_ai_content(ai_response)
    ai_paragraphs = [p.strip() for p in cleaned_ai_response.split("\n") if p.strip()]

    # 2. 关键：按知乎文章方法，统一设置文本格式（中文字体+不加粗+字号）
    def add_uniform_text(paragraph, text):
        """知乎方案：先设英文字体，再通过qn设置中文字体"""
        run = paragraph.add_run(text)
        # 步骤1：设置英文字体（必须先设置，否则中文字体可能失效）
        run.font.name = target_font
        # 步骤2：核心！通过qn函数设置中文字体（w:eastAsia属性）
        run.element.rPr.rFonts.set(qn('w:eastAsia'), target_font)
        # 步骤3：强制不加粗+统一字号
        run.font.bold = False
        run.font.size = target_font_size
        return run

    # 3. 写入AI内容（逐个段落应用统一格式）
    for para_text in ai_paragraphs:
        para = doc.add_paragraph()
        add_uniform_text(para, para_text)
        # 统一行间距和段落间距
        para.line_spacing = target_line_spacing
        para.space_after = Pt(5)

    # 保存到字节流
    doc_stream = io.BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream

# 主路由（不变）
@app.route('/', methods=['GET', 'POST'])
def index():
    font_sizes = generate_font_sizes()
    client_ip = request.remote_addr

    if request.method == 'POST':
        textbook_title = request.form.get('title', '').strip()
        selected_options = request.form.getlist('options')
        word_count = request.form.get('word_count', '').strip()
        font_size = request.form.get('font_size', '')
        line_height = request.form.get('line_height', '').strip()

        print("\n" + "=" * 50)
        print("📚 课文背景查询 - 用户输入信息")
        print("=" * 50)
        print(f"客户端IP：{client_ip}")
        print(f"课文标题：{textbook_title if textbook_title else '未输入'}")
        print(f"查询选项：{', '.join(selected_options) if selected_options else '未选择'}")
        print(f"总字数：{word_count if word_count else '未输入'}")
        print(f"字号：{font_size + '号' if font_size else '未选择'}")
        print(f"行间距：{line_height if line_height else '未输入'}")
        print("=" * 50 + "\n")

        if not (textbook_title and word_count and font_size and line_height and selected_options):
            return render_template('index.html', font_sizes=font_sizes, unwhole=True, success=False, wait=False)

        if client_ip in wait:
            return render_template('index.html', font_sizes=font_sizes, wait=True, success=False, unwhole=False)

        wait.append(client_ip)
        ai_input = f"请提供课文《{textbook_title}》的"
        if "作者简介" in selected_options:
            ai_input += "作者简介，"
        if "写作背景" in selected_options:
            ai_input += "写作背景，"
        ai_input += f"要求总字数约为{word_count}字，语言通俗易懂，结构清晰，不要包含任何格式符号、序号、特殊字符。"
        print(f"🤖 AI输入：{ai_input}")

        ai_response = ask_AI(ai_input, client_ip)
        print(f"🤖 AI回复：{ai_response}")

        # 双重清理
        ai_response = re.sub(r"[^\u4e00-\u9fa5a-zA-Z0-9\s，。！？；：]", "", ai_response)

        doc_stream = generate_word_doc(
            textbook_title=textbook_title,
            selected_options=selected_options,
            word_count=word_count,
            font_size=font_size,
            line_height=line_height,
            ai_response=ai_response
        )

        safe_title = textbook_title.replace('/', '_').replace('\\', '_')
        filename = f"课文背景_{safe_title}_{int(time.time())}.docx"
        response = make_response(send_file(
            doc_stream,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=quote(filename)
        ))
        response.headers[
            'Content-Disposition'] = f'attachment; filename="{quote(filename)}"; filename*=UTF-8\'\'{quote(filename)}'
        return response

    return render_template('index.html', font_sizes=font_sizes, success=False, unwhole=False, wait=False)

if __name__ == '__main__':
    # 部署时用0.0.0.0，端口从环境变量获取（本地运行仍用5000）
    port = int(os.getenv("PORT", 5000))

    app.run(host='0.0.0.0', port=port, debug=False)  # 部署时关闭debug
